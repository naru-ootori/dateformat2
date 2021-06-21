# coding=utf-8
'''Date format converter form US style Kibana logs to Russian docx files.'''

from pathlib import Path
from subprocess import Popen
import sys
import datetime
import binascii

from docx.enum.text import  WD_ALIGN_PARAGRAPH, WD_LINE_SPACING \
                            # pylint: disable=E0611

from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, \
                            WD_ROW_HEIGHT_RULE # pylint: disable=E0611

from PyQt5.QtWidgets import *
from PyQt5.QtGui import *
from docx import Document
from docx.shared import Cm, Pt

with open('config.txt', 'r', encoding="utf-8") as f:

    conf         = f.readlines()
    default_path = conf[0].replace('\n', '')
    word_path    = conf[1].replace('\n', '')
    start_date   = conf[2].replace('\n', '')
    end_date     = datetime.datetime.now().strftime('%d.%m.%Y')

with open('template.txt', 'r', encoding="utf-8") as template:

    template_text = template.readlines()
    paragraph_0   = template_text[0].replace('\n', '')
    paragraph_1   = template_text[1].replace('\n', '')
    paragraph_2   = template_text[2].replace('\n', '')


class MainWindow(QMainWindow):
    '''Draw main window'''

    def __init__(self):

        super().__init__()
        self.init_ui()

    def init_ui(self):
        '''UI initialization'''

        self.resize(800, 695)
        self.center()
        self.setWindowTitle('Формататор логов 2.2')

        label_output = QLabel(self)
        label_output.move(10, 10)
        label_output.setFont(QFont('Arial', 11))
        label_output.setText('Сохранить в:')

        self.qle_output = QLineEdit(self)
        self.qle_output.setGeometry(120, 10, 600, 30)
        self.qle_output.setFont(QFont('Arial', 10))
        self.qle_output.setText(default_path)

        self.button_output = QPushButton('Обзор', self)
        self.button_output.setFont(QFont('Arial', 11))
        self.button_output.setGeometry(730, 10, 60, 30)
        self.button_output.clicked.connect(self.browse_for_output)

        label_orgname = QLabel(self)
        label_orgname.move(10, 50)
        label_orgname.setFont(QFont('Arial', 11))
        label_orgname.setText('Наименование:')

        self.qle_orgname = QLineEdit(self)
        self.qle_orgname.setGeometry(120, 50, 260, 30)
        self.qle_orgname.setFont(QFont('Arial', 10))

        label_inn = QLabel(self)
        label_inn.move(390, 50)
        label_inn.setFont(QFont('Arial', 11))
        label_inn.setText('ИНН:')

        self.qle_inn = QLineEdit(self)
        self.qle_inn.setGeometry(430, 50, 200, 30)
        self.qle_inn.setFont(QFont('Arial', 10))

        self.button_savefile = QPushButton('Сформировать', self)
        self.button_savefile.setFont(QFont('Arial', 11))
        self.button_savefile.setGeometry(640, 50, 150, 30)
        self.button_savefile.clicked.connect(self.save_document)

        label_log_paste = QLabel(self)
        label_log_paste.move(10, 90)
        label_log_paste.setFont(QFont('Arial', 11))
        label_log_paste.setText('Текст лога:')

        self.progress = QProgressBar(self)
        self.progress.setTextVisible(0)
        self.progress.setGeometry(120, 90, 670, 20)
        self.progress.setMinimum(0)
        self.progress.setValue(0)

        self.log_paste = QTextEdit(self, acceptRichText = False)
        self.log_paste.setGeometry(10, 120, 780, 470)

        label_converter = QLabel(self)
        label_converter.setGeometry(10, 600, 780, 20)
        label_converter.setFont(QFont('Arial', 11))
        label_converter.setText('Преобразовать GUID в hex-формат:')

        label_guid = QLabel(self)
        label_guid.setGeometry(10, 630, 780, 20)
        label_guid.setFont(QFont('Arial', 11))
        label_guid.setText('GUID:')

        self.qle_guid = QLineEdit(self)
        self.qle_guid.setGeometry(60, 630, 300, 20)
        self.qle_guid.setFont(QFont('Arial', 11))
        self.qle_guid.setMaxLength(36)
        #self.qle_guid.setInputMask('HHHHHHHH-HHHH-HHHH-HHHH-HHHHHHHHHHHH')
        self.qle_guid.textChanged.connect(self.guid_to_hex)

        label_hex = QLabel(self)
        label_hex.setGeometry(10, 660, 780, 20)
        label_hex.setFont(QFont('Arial', 11))
        label_hex.setText('HEX:')

        self.qle_hex = QLineEdit(self)
        self.qle_hex.setGeometry(60, 660, 300, 20)
        self.qle_hex.setFont(QFont('Arial', 11))

        self.clear_button = QPushButton('Очистить поля', self)
        self.clear_button.setFont(QFont('Arial', 11))
        self.clear_button.setGeometry(530, 620, 260, 60)
        self.clear_button.clicked.connect(self.clear_fields)

        self.show()

    def center(self):
        '''Center main window'''

        window_parameters = self.frameGeometry()
        center_position   = QDesktopWidget().availableGeometry().center()
        window_parameters.moveCenter(center_position)
        self.move(window_parameters.topLeft())

    def browse_for_output(self):
        '''Open brows for output folder window'''

        output_folder = QFileDialog.getExistingDirectory(self,
                                              'Выберите папку для сохранения:')
        output_folder = str(Path(output_folder))
        self.qle_output.setText(output_folder)

    def convert_log(self):
        '''Convert log to Russian date-time format.
        Black magic fuckery happens here, hard.'''

        raw_text   = self.log_paste.toPlainText()
        split_text = raw_text.split('\n')

        final_log = []

        for raw_line in split_text:

            if len(raw_line) < 2:
                continue

            clean_line    = raw_line.lstrip('\t ').split('\t')
            clean_line[0] = clean_line[0][:-4]

            timestamp    = clean_line[0].split(' ')
            timestamp[1] = timestamp[1].translate(str.maketrans('', '',
                                                                'stndrdth'))

            if len(timestamp[1]) == 1:
                timestamp[1] = '0' + timestamp[1]

            timestamp[2] = timestamp[2].rstrip(',')

            ip_address = clean_line[1]
            if ip_address == ' - ':
                continue

            raw_datetime =(timestamp[0] + ' ' +
                           timestamp[1] + ' ' +
                           timestamp[2] + ' ' +
                           timestamp[3])

            std_datetime = datetime.datetime.strptime(raw_datetime,
                                                      '%B %d %Y %H:%M:%S')
            rus_datetime = datetime.datetime.strftime(std_datetime,
                                                      '%d.%m.%Y, %H:%M:%S')

            final_line  = [rus_datetime, ip_address]

            if final_log != [] and final_line == final_log[-1]:
                continue

            final_log.append(final_line)

        return final_log

    def table_format(self, table):
        '''Apply styles to various table elements'''

        table.style            = 'Table Grid'
        table.alignment        = WD_TABLE_ALIGNMENT.CENTER
        table.rows.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[0].height   = Cm(0.8)

        for col in table.columns:

            for cell in col.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                for par in cell.paragraphs:
                    par.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.LEFT
                    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    par.paragraph_format.space_before      = Pt(0)
                    par.paragraph_format.space_after       = Pt(0)

                    for run in par.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)

        for cell in table.row_cells(0):

            for par in cell.paragraphs:
                par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                for run in par.runs:
                    run.font.bold = True

    def save_document(self):
        '''Make and save a docx file'''

        start = datetime.datetime.now()

        source_log = self.convert_log()


        savepath   = Path(self.qle_output.text().strip())
        orgname    = self.qle_orgname.text().translate(str.maketrans('', '',
                                                       '.«»\'\"')).strip()

        if orgname == '':
            orgname = 'ООО ОРГАНИЗАЦИЯ'

        orgtype    = orgname.partition(' ')[0]
        clrname    = orgname.partition(' ')[2]

        if clrname == '':
            clrname = orgname
            orgtype = ''

        if orgtype != 'ИП':
            stdname = orgtype + ' ' + '«' + clrname + '»'
        else:
            stdname = orgname

        inn = self.qle_inn.text().strip()
        if inn == '':
            inn = '0000000000'

        table_size = len(source_log)

        self.progress.setMaximum(table_size)
        self.progress.setValue(0)

        document = Document('template.docx')

        document.sections[-1].top_margin    = Cm(1)
        document.sections[-1].bottom_margin = Cm(1)
        document.sections[-1].left_margin   = Cm(2)
        document.sections[-1].right_margin  = Cm(2)

        document.paragraphs[0].text = (paragraph_0.format(stdname, inn))
        document.add_paragraph(paragraph_1)
        document.add_paragraph(paragraph_2.format(start_date, end_date))

        for par in document.paragraphs:
            par.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
            par.paragraph_format.first_line_indent = Cm(1)
            par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            par.paragraph_format.space_before      = Pt(0)
            par.paragraph_format.space_after       = Pt(0)
            for run in par.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)

        columns           = 2
        table             = document.add_table(rows = table_size+1,
                                               cols = columns)
        hdr_cells         = table.rows[0].cells
        hdr_cells[0].text = 'Дата и время начала сеанса'
        hdr_cells[1].text = 'Интернет-адрес рабочего места абонента'

        # pylint: disable=W0212
        table_cells = table._cells

        for i in range(0, table_size):

            row_cells = table_cells[(i+1)*columns:(i+2)*columns]
            row_cells[0].text = source_log[i][0]
            row_cells[1].text = source_log[i][1]

            self.progress.setValue(i+1)

        self.table_format(table)

        docpath = str(savepath) + '\\' + '{0}.docx'.format(orgname)

        print('----')
        print('Наименование организации:', stdname)
        print('ИНН:', inn)
        print('Строк записано:', table_size)
        print('Сохранение файла:', docpath)

        try:
            document.save(docpath)
        except OSError as error_message:
            print("Error: {0} - {1}.".format(error_message.filename,
                                             error_message.strerror))

        args = [word_path, '/n', docpath]
        Popen(args)

        end = datetime.datetime.now()

        print('Затрачено времени: {0}'.format(end - start))
        print('----')

    def guid_to_hex(self):
        '''Convert GUID to hex format'''

        guid = self.qle_guid.text()

        if len(guid) != 36:
            self.qle_hex.setText('—')

        else:
            guid_map = binascii.unhexlify(guid.translate(str.maketrans('', '',
                                                                     '-\r\n ')))
            hexvalue = ''.join(map(bytes.decode, map(
                        binascii.hexlify,(guid_map[3::-1],
                                          guid_map[5:3:-1],
                                          guid_map[7:5:-1],
                                          guid_map[8:]))))
            self.qle_hex.setText(hexvalue)

    def clear_fields(self):
        '''Clear input fields'''

        self.qle_guid.setText('')
        self.log_paste.setText('')
        self.qle_inn.setText('')
        self.qle_orgname.setText('')

APP = QApplication(sys.argv)
MW = MainWindow()
APP.exec_()
